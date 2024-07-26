import os
from openpyxl import Workbook
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

start = 2


def xlsx_case(foo_file_xlsx, data_list):
    foo_file_xlsx_new = ""
    if os.path.exists(f"{foo_file_xlsx[4]}/{foo_file_xlsx[5]}.xlsx"):  # Checar asignación de nombre.
        base, extension = os.path.splitext(foo_file_xlsx[5])
        count = 0
        new_filename = f"{base}.xlsx"
        while os.path.exists(f"{foo_file_xlsx[4]}/{new_filename}"):
            count += 1
            new_filename = f"{base}_{count}.xlsx"
        foo_file_xlsx_new = f"{foo_file_xlsx[4]}/{new_filename}"
    else:
        foo_file_xlsx_new = f"{foo_file_xlsx[4]}/{foo_file_xlsx[5]}.xlsx"

    wb = Workbook()
    ws = wb.active

    for i, k in enumerate(data_list, start=start):
        [ws.cell(row=i, column=j, value=l) for j, l in enumerate(k, start=1)]

    ws["A1"] = "Original"
    ws["B1"] = "Translated"
    wb.save(foo_file_xlsx_new)
    wb.close()


def word_case(foo_file_docx, foo_text_list_docx,
              foo_font="Arial", foo_points=14, foo_color=(0, 0, 0)):
    foo_file_docx_new = ""
    if os.path.exists(f"{foo_file_docx[4]}/{foo_file_docx[5]}.docx"):  # Checar asignación de nombre.
        base, extension = os.path.splitext(foo_file_docx[5])
        count = 0
        new_filename = f"{base}.docx"
        while os.path.exists(f"{foo_file_docx[4]}/{new_filename}"):
            count += 1
            new_filename = f"{base}_{count}.docx"
        foo_file_docx_new = f"{foo_file_docx[4]}/{new_filename}"
    else:
        foo_file_docx_new = f"{foo_file_docx[4]}/{foo_file_docx[5]}.docx"
    doc = Document()
    for _ in range(len(foo_text_list_docx)):
        for __ in range(len(foo_text_list_docx[_])):
            if __ == 0:  # Checar si funciona.
                paragraph = doc.add_paragraph(f"Original: {foo_text_list_docx[_][__]}")
            else:
                paragraph = doc.add_paragraph(f"Traducción: {foo_text_list_docx[_][__]}")
            run = paragraph.runs[0]
            run.font.name = foo_font
            run.font.size = Pt(foo_points)
            run.font.color.rgb = RGBColor(*foo_color)
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        paragraph1 = doc.add_paragraph("---")
        run = paragraph1.runs[0]
        run.font.name = foo_font
        run.font.size = Pt(foo_points)
        run.font.color.rgb = RGBColor(*foo_color)
        paragraph1.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    doc.save(foo_file_docx_new)
